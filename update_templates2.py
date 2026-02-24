from docx import Document
from docx.shared import Inches
import os

def update_template(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    print(f"Processing {filepath}...")
    
    # Reload from a fresh backup if exist or create one
    backup_path = filepath + ".bak"
    if not os.path.exists(backup_path):
        import shutil
        shutil.copy2(filepath, backup_path)
    else:
        # Restore backup to process fresh
        import shutil
        shutil.copy2(backup_path, filepath)
        
    doc = Document(filepath)
    
    found_sign = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Insp. Vikram Singh" in text or "Inspector Vikram Singh" in text:
            print("Found signature line:", text)
            new_p = p.insert_paragraph_before()
            run = new_p.add_run()
            run.add_picture('vikram_sign.png', width=Inches(1.5))
            new_p.alignment = p.alignment
            new_p.paragraph_format.left_indent = p.paragraph_format.left_indent
            new_p.paragraph_format.first_line_indent = p.paragraph_format.first_line_indent
            found_sign = True
            break
            
    if not found_sign:
        print("Warning: Could not find 'Insp. Vikram Singh' in", filepath)
        
    # We want to insert stamp after his details.
    # Usually it's "Email ID..." or similar, or just at the end.
    stamp_added = False
    
    # Let's search from the found paragraph down for the end of block
    if found_sign:
        # Find the end of the block (empty line or end of doc)
        for j in range(i, len(doc.paragraphs)):
            text = doc.paragraphs[j].text.strip()
            # If we hit an empty line after the block, or specific known last lines like Email or "Dated : {DATE}"
            if "Email" in text or "Email ID" in text or "Dated" in text:
                # Add stamp after this paragraph
                stamp_p = doc.paragraphs[j].insert_paragraph_before() if "Dated" in text else doc.paragraphs[j].insert_paragraph_before() 
                # wait, if it's email, insert after it.
                pass
                
    # Actually just add it after the last non-empty line
    last_p = None
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            last_p = p
            break
            
    if last_p:
        print("Adding stamp after:", last_p.text)
        stamp_p = doc.add_paragraph()
        run = stamp_p.add_run()
        run.add_picture('vikra_stamp.png', width=Inches(2.5))
        stamp_p.alignment = last_p.alignment
    else:
        print("Warning: No text found to append stamp in", filepath)
        
    doc.save(filepath)
    print(f"Saved {filepath}")

templates = [
    'Airtel Template.docx',
    'JIO Template.docx',
    'VI Template.docx',
    'bank_letters/bank_layerwise_template.docx',
    'bank_letters/money_release_template.docx',
    'bank_letters/atm_template.docx',
    'bank_letters/cheque_withdrawal_template.docx',
    'bank_letters/aeps_template.docx'
]

for t in templates:
    update_template(t)

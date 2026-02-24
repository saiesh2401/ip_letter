import os
import json
import datetime
from bs4 import BeautifulSoup
from ipwhois import IPWhois
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict
import warnings

import zipfile
import pandas as pd

# Suppress warnings
warnings.filterwarnings("ignore")

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

class ISPProcessor:
    def __init__(self, output_dir="Generated_Letters", cache_file="isp_cache.json"):
        self.output_dir = output_dir
        # Cache file should live in the actual user directory, not the temp bundle, 
        # so we DO NOT use resource_path for it (unless we want read-only cache).
        self.cache_file = cache_file 
        self.isp_cache = self._load_cache()
        
        # Ensure output dir exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def process_airtel_reply(self, zip_path, password_str):
        """
        Process an Airtel Reply Zip (Nested Zips of CSVs).
        Returns: {
            'hits_df': pd.DataFrame (Combined valid rows),
            'misses': list (List of IPs/Files with no records)
        }
        """
        extract_root = os.path.join(self.output_dir, "temp_reply_extract")
        if os.path.exists(extract_root):
             import shutil
             shutil.rmtree(extract_root)
        os.makedirs(extract_root)
        
        hits = []
        misses = []
        
        try:
            # 1. Extract Master Zip
            password_bytes = password_str.encode("utf-8") if password_str else None
            with zipfile.ZipFile(zip_path, 'r') as zf:
                try:
                    zf.extractall(path=extract_root, pwd=password_bytes)
                except RuntimeError:
                     return None, "Incorrect Password"
                except Exception as e:
                     return None, f"Extraction Error: {e}"

            # 2. Iterate Extracted
            for root, _, files in os.walk(extract_root):
                for f in files:
                    if f.endswith(".zip"):
                        # Inner Zip Found (usually one per IP)
                        inner_zip_path = os.path.join(root, f)
                        inner_extract_dir = os.path.join(root, os.path.splitext(f)[0])
                        
                        # Extract Inner
                        try:
                            with zipfile.ZipFile(inner_zip_path, 'r') as zf_inner:
                                # Try with password, then without
                                try:
                                    zf_inner.extractall(path=inner_extract_dir, pwd=password_bytes)
                                except:
                                    zf_inner.extractall(path=inner_extract_dir)
                                    
                            # Scan for CSV inside inner dir
                            csv_found = False
                            for sub_root, _, sub_files in os.walk(inner_extract_dir):
                                for sub_f in sub_files:
                                    if sub_f.endswith(".csv"):
                                        csv_found = True
                                        csv_full_path = os.path.join(sub_root, sub_f)
                                        
                                        # Parse CSV
                                        # Airtel CSV: ~Line 7 is header.
                                        try:
                                            with open(csv_full_path, "rb") as ft:
                                                lines = ft.readlines()
                                            
                                            # Check for No Records
                                            # Usually data starts at line 7 (index 7) or 8
                                            file_content_str = b"".join(lines).decode('latin1')
                                            
                                            if "No Records Found" in file_content_str:
                                                misses.append(f) # Add Zip Name or IP to misses
                                            else:
                                                # Attempt to read as DataFrame
                                                # Standard header is usually at row 6 (0-indexed)
                                                # We can try to read with header=6
                                                try:
                                                    # Robust Header Finding
                                                    with open(csv_full_path, "r", encoding="latin1") as f_in:
                                                        lines = f_in.readlines()
                                                    
                                                    header_idx = -1
                                                    for i, line in enumerate(lines[:20]): # Scan first 20 lines
                                                        if "DSL_User_ID" in line:
                                                            header_idx = i
                                                            break
                                                            
                                                    if header_idx != -1:
                                                        # Use python engine for robust parsing of quotes/dates
                                                        # Read all columns as strings to prevent Arrow serialization errors
                                                        try:
                                                            df = pd.read_csv(csv_full_path, skiprows=header_idx, header=0, encoding='latin1', quotechar="'", skipinitialspace=True, on_bad_lines='skip', engine='python', dtype=str)
                                                        except:
                                                            # Fallback for older pandas or different env
                                                            df = pd.read_csv(csv_full_path, skiprows=header_idx, header=0, encoding='latin1', quotechar="'", skipinitialspace=True, engine='python', dtype=str)

                                                        # Filter footer artifacts
                                                        if 'DSL_User_ID' in df.columns:
                                                            df = df[~df['DSL_User_ID'].astype(str).str.contains("System generated", case=False, na=False)]
                                                            
                                                        if not df.empty:
                                                            df['Source_File'] = f
                                                            df['CSV_Path'] = csv_full_path
                                                            hits.append(df)
                                                except:
                                                     pass
                                        except:
                                            pass
                            
                            if not csv_found:
                                misses.append(f"{f} (No CSV)")

                        except Exception:
                            misses.append(f"{f} (Corrupt/PW)")
                            
            # 3. Combine Hits
            if hits:
                final_df = pd.concat(hits, ignore_index=True)
            else:
                final_df = pd.DataFrame()
                
            return {
                "hits_df": final_df,
                "misses": misses
            }, "Success"

        except Exception as e:
             return None, f"Processing Error: {e}"


    def process_jio_reply(self, file_path):
        """
        Process Jio Reply files in either .7z or .zip format.
        Automatically detects format and extracts accordingly.
        """
        import shutil
        import pandas as pd
        
        temp_dir = os.path.join(os.getcwd(), "temp_jio_extract")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        
        hits = []
        misses = []
        
        try:
            # 1. Detect file format and extract accordingly
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.7z':
                # Extract 7z file
                try:
                    import py7zr
                    with py7zr.SevenZipFile(file_path, mode='r') as z:
                        z.extractall(path=temp_dir)
                except Exception as e:
                    return None, f"Failed to extract 7z: {str(e)}"
            
            elif file_ext == '.zip':
                # Extract zip file
                try:
                    with zipfile.ZipFile(file_path, 'r') as z:
                        z.extractall(path=temp_dir)
                except Exception as e:
                    return None, f"Failed to extract zip: {str(e)}"
            
            else:
                return None, f"Unsupported file format: {file_ext}. Only .7z and .zip are supported."
                
            # 2. Walk and Find Data
            for root, dirs, files in os.walk(temp_dir):
                for f in files:
                    # Filter for interesting files
                    if f.lower().endswith(('.csv', '.xlsx', '.xls')):
                        file_full_path = os.path.join(root, f)
                        try:
                            # Read file
                            if f.lower().endswith('.csv'):
                                # Jio CSVs seem clean, header=0
                                # Read all columns as strings to prevent Arrow serialization errors
                                df = pd.read_csv(file_full_path, on_bad_lines='skip', encoding='latin1', low_memory=False, dtype=str)
                            else:
                                df = pd.read_excel(file_full_path, dtype=str)
                                
                            if not df.empty:
                                df['Source_File'] = f
                                df['CSV_Path'] = file_full_path
                                hits.append(df)
                            else:
                                misses.append(f)
                        except Exception as e:
                            # print(f"Error reading {f}: {e}")
                            misses.append(f)
                            
        except Exception as e:
            return None, f"Processing Error: {str(e)}"
            
        # Combine
        if hits:
            final_df = pd.concat(hits, ignore_index=True)
        else:
            final_df = pd.DataFrame()
            
        return {
            "hits_df": final_df,
            "misses": misses
        }, "Success"


    def process_vi_reply(self, file_path):
        """
        Process a VI (Vodafone Idea) Reply Zip file containing CSV data.
        Returns: {
            'total_rows': int,
            'output_file': str (path to combined Excel),
            'sample_data': DataFrame (first 100 rows),
            'columns': list,
            'misses': list
        }
        """
        import tempfile
        import shutil
        
        temp_dir = os.path.join(self.output_dir, "temp_vi_extract")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        
        # Create temporary Excel output file
        temp_output = tempfile.NamedTemporaryFile(mode='wb', delete=False, suffix='.xlsx', dir=self.output_dir)
        output_path = temp_output.name
        temp_output.close()
        
        total_rows = 0
        misses = []
        sample_data = None
        columns = []
        all_chunks = []
        
        try:
            # 1. Extract Zip (no password for VI)
            try:
                with zipfile.ZipFile(file_path, 'r') as zf:
                    zf.extractall(path=temp_dir)
            except Exception as e:
                return None, f"Failed to extract zip: {str(e)}"
            
            # 2. Find and process CSV files
            for root, dirs, files in os.walk(temp_dir):
                for f in files:
                    if f.lower().endswith('.csv'):
                        file_full_path = os.path.join(root, f)
                        try:
                            # VI CSV has metadata header (11 lines) before column headers
                            # Line 1: "VIL Call Data Records"
                            # Lines 2-10: Metadata (PRIVIP, Report Type, dates, etc.)
                            # Line 11: Blank
                            # Line 12: Column headers
                            # Lines 13+: Data
                            
                            # Process in chunks
                            chunk_iter = pd.read_csv(
                                file_full_path,
                                skiprows=11,  # Skip metadata header
                                chunksize=10000,
                                on_bad_lines='skip',
                                encoding='utf-8',
                                low_memory=False,
                                dtype=str
                            )
                            
                            file_has_data = False
                            for chunk in chunk_iter:
                                if not chunk.empty:
                                    file_has_data = True
                                    chunk['Source_File'] = f
                                    chunk['CSV_Path'] = file_full_path
                                    
                                    # Save sample from first chunk
                                    if sample_data is None:
                                        sample_data = chunk.head(100).copy()
                                        columns = list(chunk.columns)
                                    
                                    # Collect chunk for Excel writing
                                    all_chunks.append(chunk)
                                    total_rows += len(chunk)
                                    
                                    # Free memory every 50k rows
                                    if total_rows % 50000 == 0:
                                        import gc
                                        gc.collect()
                            
                            if not file_has_data:
                                misses.append(f)
                                
                        except Exception as e:
                            misses.append(f"{f} (Error: {str(e)})")
            
            # Write all chunks to Excel file
            if all_chunks:
                combined_df = pd.concat(all_chunks, ignore_index=True)
                combined_df.to_excel(output_path, index=False, engine='openpyxl')
                del combined_df
                import gc
                gc.collect()
            
        except Exception as e:
            return None, f"Processing Error: {str(e)}"
        
        # Return summary
        return {
            "total_rows": total_rows,
            "output_file": output_path,
            "sample_data": sample_data,
            "columns": columns,
            "misses": misses
        }, "Success"



    def _load_cache(self):
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def _save_cache(self):
        try:
            with open(self.cache_file, 'w') as f:
                json.dump(self.isp_cache, f)
        except Exception as e:
            print(f"Failed to save cache: {e}")

    def get_isp(self, ip, log_callback=None):
        if ip in self.isp_cache:
            return self.isp_cache[ip]
        
        try:
            obj = IPWhois(ip)
            res = obj.lookup_rdap()
            org = res.get('network', {}).get('name', 'Unknown')
            if not org:
                org = res.get('asn_description', 'Unknown')
            
            if log_callback:
                log_callback(f"Lookup: {ip} -> {org}")
            
            org_upper = org.upper()
            if "JIO" in org_upper or "RELIANCE" in org_upper:
                isp = "JIO"
            elif "AIRTEL" in org_upper or "BHARTI" in org_upper:
                isp = "AIRTEL"
            elif "VODAFONE" in org_upper or "VI" in org_upper or "IDEA" in org_upper:
                isp = "VI"
            elif "BSNL" in org_upper:
                isp = "BSNL"
            else:
                isp = "OTHER"
                
            self.isp_cache[ip] = isp
            self._save_cache()
            return isp
        except Exception as e:
            if log_callback:
                log_callback(f"Error looking up {ip}: {e}")
            return "Unknown"

    def parse_html(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"HTML file not found: {file_path}")

        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')

        data = []
        metadata = {"name": "Unknown", "email": "Unknown"}
        
        # Extract Metadata (Name, Email)
        # Usually in the first few sections
        list_items = soup.find_all('li')
        for li in list_items:
            text = li.get_text().strip()
            if text.startswith("Name:"):
                metadata["name"] = text.replace("Name:", "").strip()
            elif text.startswith("e-Mail:"):
                metadata["email"] = text.replace("e-Mail:", "").strip()

        tables = soup.find_all('table')
        for table in tables:
            headers = [th.get_text().strip() for th in table.find_all('th')]
            if "IP Address" in headers and "Timestamp" in headers:
                rows = table.find_all('tr')[1:] 
                for row in rows:
                    cols = row.find_all('td')
                    if len(cols) >= 2:
                        ts = cols[0].get_text().strip()
                        ip = cols[1].get_text().strip()
                        if ip:
                             data.append({'timestamp': ts, 'ip': ip})
        return metadata, data

    def process_data(self, raw_data, progress_callback=None, log_callback=None):
        grouped = defaultdict(list)
        total = len(raw_data)
        
        for idx, entry in enumerate(raw_data):
            if progress_callback:
                progress_callback(idx + 1, total)
                
            ip = entry['ip']
            ts_str = entry['timestamp']
            
            try:
                # Format is "YYYY-MM-DD HH:MM:SS Z", which is UTC.
                # Remove Z, parse time, and add 5h 30m for IST.
                dt_utc = datetime.datetime.strptime(ts_str.replace(" Z", ""), "%Y-%m-%d %H:%M:%S")
                dt = dt_utc + datetime.timedelta(hours=5, minutes=30)
            except:
                dt = datetime.datetime.now()
                
            isp = self.get_isp(ip, log_callback)
            entry['datetime'] = dt
            grouped[isp].append(entry)
            
        return grouped

    def fill_jio_excel(self, entries, template_path, output_name="JIO_Data.xlsx"):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
            return False, "Template missing"

        wb = openpyxl.load_workbook(template_path)
        sheet = wb.active
        current_row = sheet.max_row + 1
        
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            from_dt = dt - datetime.timedelta(minutes=5)
            to_dt = dt + datetime.timedelta(minutes=5)
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=from_dt.strftime("%Y%m%d"))
            sheet.cell(row=current_row, column=4, value=from_dt.strftime("%H%M%S"))
            sheet.cell(row=current_row, column=5, value=to_dt.strftime("%Y%m%d"))
            sheet.cell(row=current_row, column=6, value=to_dt.strftime("%H%M%S"))
            current_row += 1
            
        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_jio_txt(self, entries, output_name="JIO_Data.txt"):
        out_path = os.path.join(self.output_dir, output_name)
        
        # Header from sample file
        # Note: Sample had "From Time" twice. I will keep it exactly as sample unless directed otherwise,
        # but standard logic usually suggests To Time. I'll stick to the sample header text to be safe.
        header = "Type\tSearch Value\tFrom Date YYYYMMDD\tFrom Time HHMMSS (IST)\tTo Date YYYYMMDD\tFrom Time HHMMSS (IST)\n"
        
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(header)
                
                for item in entries:
                    ip = item['ip']
                    dt = item['datetime']
                    from_dt = dt - datetime.timedelta(minutes=5)
                    to_dt = dt + datetime.timedelta(minutes=5)
                    ip_type = "IPV6" if ":" in ip else "IPV4"
                    
                    # JIO Format: YYYYMMDD, HHMMSS
                    row = [
                        ip_type,
                        ip,
                        from_dt.strftime("%Y%m%d"),
                        from_dt.strftime("%H%M%S"),
                        to_dt.strftime("%Y%m%d"),
                        to_dt.strftime("%H%M%S")
                    ]
                    f.write("\t".join(row) + "\n")
                    
            return True, out_path
        except Exception as e:
            return False, str(e)

    def fill_airtel_excel(self, entries, template_path, output_name="Airtel_Data.xlsx"):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
             return False, "Template missing"

        wb = openpyxl.load_workbook(template_path)
        sheet = wb.active
        current_row = sheet.max_row + 1
        
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=dt.strftime("%d-%b-%Y"))
            sheet.cell(row=current_row, column=4, value=dt.strftime("%H:%M:%S"))
            current_row += 1

        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_generic_excel(self, entries, output_name="ISP_Data.xlsx"):
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = "IP Data"
        
        # Headers
        headers = ["IP Type", "IP Address", "Date", "Time", "From Date", "To Date"]
        for col_num, header in enumerate(headers, 1):
            sheet.cell(row=1, column=col_num, value=header)
            
        current_row = 2
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            from_dt = dt - datetime.timedelta(minutes=5)
            to_dt = dt + datetime.timedelta(minutes=5)
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=dt.strftime("%d-%m-%Y"))
            sheet.cell(row=current_row, column=4, value=dt.strftime("%H:%M:%S"))
            sheet.cell(row=current_row, column=5, value=from_dt.strftime("%d-%m-%Y %H:%M:%S"))
            sheet.cell(row=current_row, column=6, value=to_dt.strftime("%d-%m-%Y %H:%M:%S"))
            current_row += 1
            
        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_word_letter(self, entries, isp_name, template_path, output_name=None, metadata=None):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
            return False, "Template missing"

        if output_name is None:
            output_name = f"{isp_name}_Request_Letter.docx"

        doc = Document(template_path)
        
        # Map simple ISP code to Full Name for the placeholder
        isp_full_names = {
            "JIO": "Reliance Jio Infocomm Ltd.",
            "AIRTEL": "Bharti Airtel Ltd.",
            "VI": "Vodafone Idea Ltd."
        }
        full_isp_name = isp_full_names.get(isp_name, isp_name)

        # Replace Placeholders in Paragraphs
        if metadata:
            name = metadata.get("name", "")
            email = metadata.get("email", "")
            fir_no = metadata.get("fir_no", "")
            fir_date = metadata.get("fir_date", "")
            current_date = datetime.datetime.now().strftime("%d.%m.%Y")
            
            replacements = {
                "{NAME}": name,
                "{EMAIL}": email,
                "{FIR_NO}": fir_no,
                "{FIR_DATE}": fir_date,
                "{ISP_NAME}": full_isp_name,
                "{DATE}": current_date
            }

            for p in doc.paragraphs:
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(val))
        
        target_table = None
        for table in doc.tables:
            if not table.rows: continue
            headers = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            if any("Search Value" in h for h in headers) or any("IP" in h for h in headers):
                target_table = table
                break
                
        if target_table:
            # Store the header row's border formatting only (not shading)
            from docx.oxml import parse_xml
            from copy import deepcopy
            
            header_row = target_table.rows[0]
            # Store ONLY border formatting from header cells (not shading)
            header_cell_borders = []
            for cell in header_row.cells:
                # Get the tcPr (table cell properties) element
                tc_pr = cell._element.tcPr
                if tc_pr is not None:
                    # Extract only the tcBorders element (not shading)
                    tc_borders = tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                    if tc_borders is not None:
                        header_cell_borders.append(deepcopy(tc_borders))
                    else:
                        header_cell_borders.append(None)
                else:
                    header_cell_borders.append(None)
            
            # CLEAR SAMPLE DATA: Remove all rows except the header
            # Iterate backwards to avoid index issues
            for i in range(len(target_table.rows) - 1, 0, -1):
                row = target_table.rows[i]
                # Remove XML element
                row._element.getparent().remove(row._element)

            # Now populate with new data
            for item in entries:
                ip = item['ip']
                dt = item['datetime']
                from_dt = dt - datetime.timedelta(minutes=5)
                to_dt = dt + datetime.timedelta(minutes=5)
                ip_type = "IPV6" if ":" in ip else "IPV4"

                # Create new row
                row = target_table.add_row()
                
                if len(row.cells) < len(target_table.columns):
                     # Should rarely happen with add_row() on uniform table
                     pass
                
                cells = row.cells
                
                # Define Format based on ISP
                if isp_name == "JIO":
                    d_fmt = "%Y%m%d"
                    t_fmt = "%H%M%S"
                elif isp_name == "VI":
                    # Image showed dots: 22.09.2025
                    d_fmt = "%d.%m.%Y" 
                    t_fmt = "%H:%M:%S"
                elif isp_name == "AIRTEL":
                    # User requested: 24/Jan/2025
                    d_fmt = "%d/%b/%Y"
                    t_fmt = "%H:%M:%S"
                else:
                    d_fmt = "%d-%m-%Y"
                    t_fmt = "%H:%M:%S"

                fmt_combined = f"{d_fmt}\n{t_fmt}"
                
                if len(cells) >= 6:
                    cells[0].text = ip_type
                    cells[1].text = ip
                    cells[2].text = from_dt.strftime(d_fmt) 
                    cells[3].text = from_dt.strftime(t_fmt)
                    cells[4].text = to_dt.strftime(d_fmt)
                    cells[5].text = to_dt.strftime(t_fmt)
                elif len(cells) >= 4:
                     cells[0].text = ip_type
                     cells[1].text = ip
                     cells[2].text = from_dt.strftime(fmt_combined)
                     cells[3].text = to_dt.strftime(fmt_combined)
                elif len(cells) == 3:
                     # 3 column table: IP, From Date Time, To Date Time
                     fmt_space = f"{d_fmt} {t_fmt}"
                     cells[0].text = ip
                     cells[1].text = from_dt.strftime(fmt_space)
                     cells[2].text = to_dt.strftime(fmt_space)
                # Apply border formatting from header to each cell
                for idx, cell in enumerate(cells):
                    if idx < len(header_cell_borders) and header_cell_borders[idx] is not None:
                        # Ensure tcPr exists
                        tc_pr = cell._element.tcPr
                        if tc_pr is None:
                            # Create tcPr if it doesn't exist
                            tc_pr = parse_xml('<w:tcPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                            cell._element.insert(0, tc_pr)
                        
                        # Remove existing borders if any
                        existing_borders = tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                        if existing_borders is not None:
                            tc_pr.remove(existing_borders)
                        
                        # Add the copied borders from header
                        tc_pr.append(deepcopy(header_cell_borders[idx]))
                
            out_path = os.path.join(self.output_dir, output_name)
            doc.save(out_path)
            return True, out_path
        else:
            return False, "Table not found"


class BankLetterProcessor:
    """
    Processes bank transaction Excel files and generates appropriate letters
    for different transaction types (Money Transfer, On Hold, ATM, Cheque).
    """
    
    VERSION = "2.1.0"  # Version with custom_subject and custom_message support
    
    def __init__(self, output_dir="Generated_Letters"):
        self.output_dir = output_dir
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def parse_bank_excel(self, file_path):
        """
        Parse the bank transaction Excel file and return all relevant sheets.
        Returns: dict with sheet names as keys and DataFrames as values
        """
        try:
            sheets = {
                'Money Transfer to': pd.read_excel(file_path, sheet_name='Money Transfer to'),
                'Transaction put on hold': pd.read_excel(file_path, sheet_name='Transaction put on hold'),
                'Withdrawal through ATM': pd.read_excel(file_path, sheet_name='Withdrawal through ATM'),
                'Cash Withdrawal through Cheque': pd.read_excel(file_path, sheet_name='Cash Withdrawal through Cheque'),
                'AEPS': pd.read_excel(file_path, sheet_name='AEPS')
            }
            return sheets, None
        except Exception as e:
            return None, f"Error parsing Excel: {str(e)}"
    
    def get_available_layers(self, df):
        """
        Get the maximum layer number available in the Money Transfer sheet.
        Returns: int (max layer number)
        """
        if 'Layer' in df.columns:
            # Filter out NaN and get unique layers
            layers = df['Layer'].dropna().unique()
            if len(layers) > 0:
                return int(max(layers))
        return 0
    
    def group_by_bank(self, df, bank_column='Action Taken By bank'):
        """
        Group transactions by bank name.
        Returns: dict with bank names as keys and list of transaction rows as values
        """
        grouped = defaultdict(list)
        
        for idx, row in df.iterrows():
            bank_name = row.get(bank_column, 'Unknown Bank')
            if pd.notna(bank_name) and bank_name != '':
                grouped[str(bank_name)].append(row)
        
        return grouped
    
    def set_table_borders(self, table):
        """
        Apply borders to all cells in a table for better visibility.
        """
        tbl = table._tbl
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                
                # Create border elements
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # Border size
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')  # Black color
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    
    def sanitize_filename(self, name):
        """Clean bank name for use in filename"""
        # Remove special characters and replace spaces with underscores
        import re
        name = re.sub(r'[^\w\s-]', '', name)
        name = re.sub(r'[\s]+', '_', name)
        return name[:100]  # Limit length
        
    def _convert_to_pdf(self, docx_path, output_dir):
        """
        Convert a DOCX file to PDF using LibreOffice in headless mode.
        Returns True if successful, False otherwise.
        """
        try:
            import subprocess
            import platform
            
            # Determine correct soffice command based on OS
            sys_os = platform.system()
            if sys_os == 'Darwin':  # macOS
                soffice_cmd = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
            elif sys_os == 'Windows':
                # Typical windows paths
                soffice_cmd_options = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                soffice_cmd = 'soffice' # fallback
                for cmd in soffice_cmd_options:
                    if os.path.exists(cmd):
                        soffice_cmd = cmd
                        break
            else: # Linux / Streamlit Cloud
                soffice_cmd = 'libreoffice'
                
            # Attempt to run conversion
            # --headless --convert-to pdf --outdir <dir> <file>
            result = subprocess.run([
                soffice_cmd, 
                '--headless', 
                '--convert-to', 'pdf', 
                '--outdir', output_dir, 
                docx_path
            ], capture_output=True, text=True)
            
            if result.returncode == 0:
                return True
            else:
                print(f"PDF Conversion failed: {result.stderr}")
                return False
        except Exception as e:
            print(f"Error during PDF conversion: {e}")
            return False
    
    def generate_layerwise_letters(self, df, num_layers, output_subdir="Sheet1_MoneyTransfer", custom_subject=None, custom_message=None):
        """
        Generate bank-wise letters for Money Transfer transactions (Sheet 1).
        Filters by layer number and groups by bank.
        Returns: list of generated file paths
        """
        output_dir = os.path.join(self.output_dir, output_subdir)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        generated_files = []
        
        # Filter by layers
        if 'Layer' in df.columns:
            df_filtered = df[df['Layer'] <= num_layers].copy()
        else:
            df_filtered = df.copy()
        
        # Group by bank (using Bank/FIs column for Sheet 1)
        grouped = self.group_by_bank(df_filtered, bank_column='Bank/FIs')
        
        # Generate letter for each bank
        for bank_name, transactions in grouped.items():
            if len(transactions) == 0:
                continue
            
            template_path = "bank_letters/bank_layerwise_template.docx"
            safe_bank_name = self.sanitize_filename(bank_name)
            output_name = f"{safe_bank_name}_Layer1-{num_layers}.docx"
            output_path = os.path.join(output_dir, output_name)
            
            success = self._fill_layerwise_template(
                template_path, 
                bank_name, 
                transactions, 
                output_path,
                num_layers,
                custom_subject,
                custom_message
            )
            
            if success:
                # Add DOCX
                generated_files.append({
                    'path': output_path,
                    'bank': bank_name,
                    'count': len(transactions),
                    'type': 'Money Transfer'
                })
                
                # Convert to PDF and add
                pdf_path = output_path.replace('.docx', '.pdf')
                if self._convert_to_pdf(output_path, output_dir) and os.path.exists(pdf_path):
                    generated_files.append({
                        'path': pdf_path,
                        'bank': f"{bank_name} (PDF)",
                        'count': len(transactions),
                        'type': 'Money Transfer'
                    })
        
        return generated_files
    
    def generate_money_release_letters(self, df, output_subdir="Sheet2_OnHold", custom_subject=None, custom_message=None, custom_court_order=None, custom_beneficiary=None):
        """
        Generate bank-wise letters for Transaction on Hold (Sheet 2).
        Returns: list of generated file paths
        """
        output_dir = os.path.join(self.output_dir, output_subdir)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        generated_files = []
        grouped = self.group_by_bank(df, bank_column='Action Taken By bank')
        
        for bank_name, transactions in grouped.items():
            if len(transactions) == 0:
                continue
            
            template_path = "bank_letters/money_release_template.docx"
            safe_bank_name = self.sanitize_filename(bank_name)
            output_name = f"{safe_bank_name}_OnHold.docx"
            output_path = os.path.join(output_dir, output_name)
            
            success = self._fill_money_release_template(
                template_path, 
                bank_name, 
                transactions, 
                output_path,
                custom_subject,
                custom_message,
                custom_court_order,
                custom_beneficiary
            )
            
            if success:
                # Add DOCX
                generated_files.append({
                    'path': output_path,
                    'bank': bank_name,
                    'count': len(transactions),
                    'type': 'Transaction On Hold'
                })
                
                # Convert to PDF and add
                pdf_path = output_path.replace('.docx', '.pdf')
                if self._convert_to_pdf(output_path, output_dir) and os.path.exists(pdf_path):
                    generated_files.append({
                        'path': pdf_path,
                        'bank': f"{bank_name} (PDF)",
                        'count': len(transactions),
                        'type': 'Transaction On Hold'
                    })
        
        return generated_files
    
    def generate_atm_letters(self, df, output_subdir="Sheet3_ATM", custom_subject=None, custom_message=None):
        """
        Generate bank-wise letters for ATM Withdrawals (Sheet 3).
        Returns: list of generated file paths
        """
        output_dir = os.path.join(self.output_dir, output_subdir)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        generated_files = []
        grouped = self.group_by_bank(df, bank_column='Action Taken By bank')
        
        for bank_name, transactions in grouped.items():
            if len(transactions) == 0:
                continue
            
            template_path = "bank_letters/atm_template.docx"
            safe_bank_name = self.sanitize_filename(bank_name)
            output_name = f"{safe_bank_name}_ATM.docx"
            output_path = os.path.join(output_dir, output_name)
            
            success = self._fill_atm_template(
                template_path, 
                bank_name, 
                transactions, 
                output_path,
                custom_subject,
                custom_message
            )
            
            if success:
                # Add DOCX
                generated_files.append({
                    'path': output_path,
                    'bank': bank_name,
                    'count': len(transactions),
                    'type': 'ATM Withdrawal'
                })
                
                # Convert to PDF and add
                pdf_path = output_path.replace('.docx', '.pdf')
                if self._convert_to_pdf(output_path, output_dir) and os.path.exists(pdf_path):
                    generated_files.append({
                        'path': pdf_path,
                        'bank': f"{bank_name} (PDF)",
                        'count': len(transactions),
                        'type': 'ATM Withdrawal'
                    })
        
        return generated_files
    
    def generate_cheque_letters(self, df, output_subdir="Sheet4_Cheque", custom_subject=None, custom_message=None):
        """
        Generate bank-wise letters for Cheque Withdrawals (Sheet 4).
        Returns: list of generated file paths
        """
        output_dir = os.path.join(self.output_dir, output_subdir)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        generated_files = []
        grouped = self.group_by_bank(df, bank_column='Action Taken By bank')
        
        for bank_name, transactions in grouped.items():
            if len(transactions) == 0:
                continue
            
            template_path = "bank_letters/cheque_withdrawal_template.docx"
            safe_bank_name = self.sanitize_filename(bank_name)
            output_name = f"{safe_bank_name}_Cheque.docx"
            output_path = os.path.join(output_dir, output_name)
            
            success = self._fill_cheque_template(
                template_path, 
                bank_name, 
                transactions, 
                output_path,
                custom_subject,
                custom_message
            )
            
            if success:
                # Add DOCX
                generated_files.append({
                    'path': output_path,
                    'bank': bank_name,
                    'count': len(transactions),
                    'type': 'Cheque Withdrawal'
                })
                
                # Convert to PDF and add
                pdf_path = output_path.replace('.docx', '.pdf')
                if self._convert_to_pdf(output_path, output_dir) and os.path.exists(pdf_path):
                    generated_files.append({
                        'path': pdf_path,
                        'bank': f"{bank_name} (PDF)",
                        'count': len(transactions),
                        'type': 'Cheque Withdrawal'
                    })
        
        return generated_files
    
    def _fill_layerwise_template(self, template_path, bank_name, transactions, output_path, num_layers, custom_subject=None, custom_message=None):
        """Fill the bank layerwise template with transaction data"""
        try:
            doc = Document(template_path)
            
            # Replace placeholders in the document
            for paragraph in doc.paragraphs:
                # Replace bank name
                if 'IDFC First Bank' in paragraph.text:
                    paragraph.text = paragraph.text.replace('IDFC First Bank', bank_name)
                
                # Replace custom subject if provided
                if custom_subject and '{{SUBJECT}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SUBJECT}}', custom_subject)
                
                # Replace custom message if provided
                if custom_message and '{{MESSAGE}}' in paragraph.text:
                    # Clear original paragraph text to remove placeholder without breaking XML
                    paragraph.text = paragraph.text.replace('{{MESSAGE}}', '')
                    
                    # Split message into lines and insert as new paragraphs with 0 spacing
                    lines = custom_message.split('\n')
                    for line in lines:
                        if not line.strip(): continue
                        new_p = paragraph.insert_paragraph_before(line)
                        new_p.paragraph_format.space_before = Pt(0)
                        new_p.paragraph_format.space_after = Pt(0)
                
                # Also check for runs within paragraphs for better replacement
                for run in paragraph.runs:
                    if 'IDFC First Bank' in run.text:
                        run.text = run.text.replace('IDFC First Bank', bank_name)
                    if custom_subject and '{{SUBJECT}}' in run.text:
                        run.text = run.text.replace('{{SUBJECT}}', custom_subject)

            
            # Find and populate the table
            if doc.tables:
                table = doc.tables[0]
                
                # Clear existing rows (keep header)
                for i in range(len(table.rows) - 1, 0, -1):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                
                # Deduplicate transactions by Account No, keeping the one with lowest layer
                unique_transactions = {}
                for txn in transactions:
                    account_no = txn.get('Account No', '')
                    if pd.notna(account_no):
                        # Convert to string for consistent comparison
                        acc_key = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        
                        # If this account hasn't been seen, or has a lower layer, keep it
                        if acc_key not in unique_transactions:
                            unique_transactions[acc_key] = txn
                        else:
                            # Compare layers - keep the one with lower layer number
                            existing_layer = unique_transactions[acc_key].get('Layer', float('inf'))
                            current_layer = txn.get('Layer', float('inf'))
                            
                            # Handle NaN values
                            if pd.isna(existing_layer):
                                existing_layer = float('inf')
                            if pd.isna(current_layer):
                                current_layer = float('inf')
                            
                            if current_layer < existing_layer:
                                unique_transactions[acc_key] = txn
                
                # Add unique transaction rows
                for idx, txn in enumerate(unique_transactions.values(), 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    if len(cells) >= 3:
                        cells[0].text = str(idx)
                        # Column G: Account No - format as text to prevent scientific notation
                        account_no = txn.get('Account No', '')
                        if pd.notna(account_no):
                            cells[1].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[1].text = ''
                        
                        # Column H: IFSC Code with Layer annotation
                        ifsc = str(txn.get('Ifsc Code', '')) if pd.notna(txn.get('Ifsc Code', '')) else ''
                        layer = txn.get('Layer', '')
                        if pd.notna(layer):
                            cells[2].text = f"{ifsc}       (L{int(layer)})"
                        else:
                            cells[2].text = ifsc
            
            # Apply borders to the table
            if doc.tables:
                self.set_table_borders(doc.tables[0])
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling layerwise template: {e}")
            return False
    
    def _fill_money_release_template(self, template_path, bank_name, transactions, output_path, custom_subject=None, custom_message=None, custom_court_order=None, custom_beneficiary=None):
        """Fill the money release template with transaction data"""
        try:
            doc = Document(template_path)
            
            # Replace placeholders in the document
            for paragraph in doc.paragraphs:
                # Replace bank names
                if 'IDFC Bank' in paragraph.text or 'IDFC First Bank' in paragraph.text:
                    paragraph.text = paragraph.text.replace('IDFC Bank', bank_name)
                    paragraph.text = paragraph.text.replace('IDFC First Bank', bank_name)
                
                # Replace custom subject if provided
                if custom_subject and '{{SUBJECT}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SUBJECT}}', custom_subject)
                
                # Replace custom message if provided
                if custom_message and '{{MESSAGE}}' in paragraph.text:
                    # Clear original paragraph text to remove placeholder without breaking XML
                    paragraph.text = paragraph.text.replace('{{MESSAGE}}', '')
                    
                    # Split message into lines and insert as new paragraphs with 0 spacing
                    lines = custom_message.split('\n')
                    for line in lines:
                        if not line.strip(): continue
                        new_p = paragraph.insert_paragraph_before(line)
                        new_p.paragraph_format.space_before = Pt(0)
                        new_p.paragraph_format.space_after = Pt(0)
                
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Handle COURT_ORDER and BENEFICIARY specially since they might be in the same paragraph
                if '{{COURT_ORDER}}' in paragraph.text and '{{BENEFICIARY}}' in paragraph.text:
                    if custom_court_order:
                        p_court = paragraph.insert_paragraph_before(custom_court_order)
                        p_court.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    paragraph.text = ''
                    if custom_beneficiary:
                        run = paragraph.add_run(custom_beneficiary)
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Reset any inherited indentation
                        paragraph.paragraph_format.left_indent = 0
                        paragraph.paragraph_format.first_line_indent = 0
                else:
                    if custom_court_order and '{{COURT_ORDER}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{COURT_ORDER}}', custom_court_order)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    if custom_beneficiary and '{{BENEFICIARY}}' in paragraph.text:
                        paragraph.text = ''
                        run = paragraph.add_run(custom_beneficiary)
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Reset any inherited indentation
                        paragraph.paragraph_format.left_indent = 0
                        paragraph.paragraph_format.first_line_indent = 0
                
                # Also check runs for other placeholders
                for run in paragraph.runs:
                    if 'IDFC Bank' in run.text or 'IDFC First Bank' in run.text:
                        run.text = run.text.replace('IDFC Bank', bank_name)
                        run.text = run.text.replace('IDFC First Bank', bank_name)
                    if custom_subject and '{{SUBJECT}}' in run.text:
                        run.text = run.text.replace('{{SUBJECT}}', custom_subject)

            
            # Find and populate the table
            if doc.tables:
                table = doc.tables[0]
                
                # Clear existing rows (keep header)
                for i in range(len(table.rows) - 1, 0, -1):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                
                # Add transaction rows
                for idx, txn in enumerate(transactions, 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    # Template columns: Sr. No. | Account No | Transaction Id | Money Put on hold
                    if len(cells) >= 4:
                        # Column 0: Serial Number
                        cells[0].text = str(idx)
                        
                        # Column 1: Account No - try multiple possible column names
                        account_no = (txn.get('Account No./ (Wallet /PG/PA) Id') or 
                                     txn.get('Account No') or 
                                     txn.get('Account Number') or '')
                        if pd.notna(account_no):
                            cells[1].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[1].text = ''
                        
                        # Column 2: Transaction Id
                        txn_id = (txn.get('Transaction Id / UTR Number') or 
                                 txn.get('Transaction Id') or 
                                 txn.get('UTR Number') or '')
                        cells[2].text = str(txn_id) if pd.notna(txn_id) else ''
                        
                        # Column 3: Money Put on hold
                        amount = (txn.get('Put on hold Amount') or 
                                 txn.get('Amount') or 
                                 txn.get('Money Put on hold') or '')
                        if pd.notna(amount):
                            try:
                                cells[3].text = f"{float(amount):,.2f}" if isinstance(amount, (int, float)) else str(amount)
                            except:
                                cells[3].text = str(amount)
                        else:
                            cells[3].text = ''
                    elif len(cells) >= 3:
                        # Fallback for 3-column table (no serial number column in template)
                        # Column 0: Account No
                        account_no = (txn.get('Account No./ (Wallet /PG/PA) Id') or 
                                     txn.get('Account No') or 
                                     txn.get('Account Number') or '')
                        if pd.notna(account_no):
                            cells[0].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[0].text = ''
                        
                        # Column 1: Transaction Id
                        txn_id = (txn.get('Transaction Id / UTR Number') or 
                                 txn.get('Transaction Id') or 
                                 txn.get('UTR Number') or '')
                        cells[1].text = str(txn_id) if pd.notna(txn_id) else ''
                        
                        # Column 2: Money Put on hold
                        amount = (txn.get('Put on hold Amount') or 
                                 txn.get('Amount') or 
                                 txn.get('Money Put on hold') or '')
                        if pd.notna(amount):
                            try:
                                cells[2].text = f"{float(amount):,.2f}" if isinstance(amount, (int, float)) else str(amount)
                            except:
                                cells[2].text = str(amount)
                        else:
                            cells[2].text = ''
            
            # Apply borders to the table
            if doc.tables:
                self.set_table_borders(doc.tables[0])
                
                # Remove extra empty paragraphs after the table to reduce space
                table_element = doc.tables[0]._element
                parent = table_element.getparent()
                found_table = False
                elements_to_remove = []
                for child in parent:
                    if child == table_element:
                        found_table = True
                        continue
                    if found_table and child.tag.endswith('p'):
                        texts = child.xpath('.//w:t/text()')
                        combined_text = "".join(texts).strip()
                        if not combined_text:
                            elements_to_remove.append(child)
                        else:
                            break
                            
                # Keep one empty paragraph for a single newline buffer
                if len(elements_to_remove) > 1:
                    for child in elements_to_remove[:-1]:
                        parent.remove(child)
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling money release template: {e}")
            print(f"Available columns in transaction data: {list(transactions[0].keys()) if len(transactions) > 0 else 'No transactions'}")
            import traceback
            traceback.print_exc()
            return False
    
    def _fill_atm_template(self, template_path, bank_name, transactions, output_path, custom_subject=None, custom_message=None):
        """Fill the ATM template with transaction data"""
        try:
            doc = Document(template_path)
            
            # Replace placeholders in the document
            for paragraph in doc.paragraphs:
                # Replace bank name
                if 'Bandhan Bank' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Bandhan Bank', bank_name)
                
                # Replace custom subject if provided
                if custom_subject and '{{SUBJECT}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SUBJECT}}', custom_subject)
                
                # Replace custom message if provided
                if custom_message and '{{MESSAGE}}' in paragraph.text:
                    # Clear original paragraph text to remove placeholder without breaking XML
                    paragraph.text = paragraph.text.replace('{{MESSAGE}}', '')
                    
                    # Split message into lines and insert as new paragraphs with 0 spacing
                    lines = custom_message.split('\n')
                    for line in lines:
                        if not line.strip(): continue
                        new_p = paragraph.insert_paragraph_before(line)
                        new_p.paragraph_format.space_before = Pt(0)
                        new_p.paragraph_format.space_after = Pt(0)
                
                # Also check runs
                for run in paragraph.runs:
                    if 'Bandhan Bank' in run.text:
                        run.text = run.text.replace('Bandhan Bank', bank_name)
                    if custom_subject and '{{SUBJECT}}' in run.text:
                        run.text = run.text.replace('{{SUBJECT}}', custom_subject)

            
            # Find and populate the table
            if doc.tables:
                table = doc.tables[0]
                
                # Clear existing rows (keep header)
                for i in range(len(table.rows) - 1, 0, -1):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                
                # Add transaction rows
                for idx, txn in enumerate(transactions, 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    # Template columns: Sl. No. | Bank Account No. | ATM ID & Place | Transaction Date & Time | Withdrawal Amount
                    if len(cells) >= 5:
                        cells[0].text = str(idx)
                        
                        # Column C: Account No./ (Wallet /PG/PA) Id
                        account_no = txn.get('Account No./ (Wallet /PG/PA) Id', '')
                        if pd.notna(account_no):
                            cells[1].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[1].text = ''
                        
                        # Column H: ATM ID and Column I: Place/Location of ATM (combined in one cell)
                        atm_id = str(txn.get('ATM ID', '')) if pd.notna(txn.get('ATM ID', '')) else ''
                        place = str(txn.get('Place/Location of ATM', '')) if pd.notna(txn.get('Place/Location of ATM', '')) else ''
                        cells[2].text = f"{atm_id}\n{place}" if atm_id and place else (atm_id or place)
                        
                        # Column E: Withdrawal Date & Time
                        cells[3].text = str(txn.get('Withdrawal Date & Time', '')) if pd.notna(txn.get('Withdrawal Date & Time', '')) else ''
                        
                        # Column F: Withdrawal Amount
                        amount = txn.get('Withdrawal Amount', '')
                        if pd.notna(amount):
                            # Keep original formatting if it's a string with commas
                            cells[4].text = str(amount)
                        else:
                            cells[4].text = ''
            
            # Apply borders to the table
            if doc.tables:
                self.set_table_borders(doc.tables[0])
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling ATM template: {e}")
            return False
    
    def _fill_cheque_template(self, template_path, bank_name, transactions, output_path, custom_subject=None, custom_message=None):
        """Fill the cheque withdrawal template with transaction data"""
        try:
            doc = Document(template_path)
            
            # Replace placeholders in the document
            for paragraph in doc.paragraphs:
                # Replace bank name
                if 'Punjab National Bank' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Punjab National Bank', bank_name)
                
                # Replace custom subject if provided
                if custom_subject and '{{SUBJECT}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SUBJECT}}', custom_subject)
                
                # Replace custom message if provided
                if custom_message and '{{MESSAGE}}' in paragraph.text:
                    # Clear original paragraph text to remove placeholder without breaking XML
                    paragraph.text = paragraph.text.replace('{{MESSAGE}}', '')
                    
                    # Split message into lines and insert as new paragraphs with 0 spacing
                    lines = custom_message.split('\n')
                    for line in lines:
                        if not line.strip(): continue
                        new_p = paragraph.insert_paragraph_before(line)
                        new_p.paragraph_format.space_before = Pt(0)
                        new_p.paragraph_format.space_after = Pt(0)
                
                # Also check runs
                for run in paragraph.runs:
                    if 'Punjab National Bank' in run.text:
                        run.text = run.text.replace('Punjab National Bank', bank_name)
                    if custom_subject and '{{SUBJECT}}' in run.text:
                        run.text = run.text.replace('{{SUBJECT}}', custom_subject)

            
            # Find and populate the table
            if doc.tables:
                table = doc.tables[0]
                
                # Clear existing rows (keep header)
                for i in range(len(table.rows) - 1, 0, -1):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                
                # Add transaction rows
                for idx, txn in enumerate(transactions, 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    # Template columns: Sl. No. | Bank Account No. | Cheque no. | IFSC Code | Transaction Date & Time | Withdrawal Amount
                    if len(cells) >= 6:
                        cells[0].text = str(idx)
                        
                        # Column E: Account No
                        account_no = txn.get('Account No', '')
                        if pd.notna(account_no):
                            cells[1].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[1].text = ''
                        
                        # Column G: Cheque No
                        cheque_no = txn.get('Cheque No', '')
                        if pd.notna(cheque_no):
                            cells[2].text = str(int(cheque_no)) if isinstance(cheque_no, (int, float)) else str(cheque_no)
                        else:
                            cells[2].text = ''
                        
                        # Column F: Ifsc Code
                        cells[3].text = str(txn.get('Ifsc Code', '')) if pd.notna(txn.get('Ifsc Code', '')) else ''
                        
                        # Column H: Withdrawal Date & Time
                        cells[4].text = str(txn.get('Withdrawal Date & Time', '')) if pd.notna(txn.get('Withdrawal Date & Time', '')) else ''
                        
                        # Column I: Withdrawal Amount
                        amount = txn.get('Withdrawal Amount', '')
                        if pd.notna(amount):
                            cells[5].text = str(amount)
                        else:
                            cells[5].text = ''
            
            # Apply borders to the table
            if doc.tables:
                self.set_table_borders(doc.tables[0])
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling cheque template: {e}")
            return False

    def generate_aeps_letters(self, df, output_subdir="Sheet5_AEPS", custom_subject=None, custom_message=None):
        """
        Generate bank-wise letters for AEPS Withdrawals (AEPS sheet).
        Returns: list of generated file paths
        """
        output_dir = os.path.join(self.output_dir, output_subdir)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        generated_files = []
        grouped = self.group_by_bank(df, bank_column='Action Taken By bank')
        
        for bank_name, transactions in grouped.items():
            if len(transactions) == 0:
                continue
            
            template_path = "bank_letters/aeps_template.docx"
            safe_bank_name = self.sanitize_filename(bank_name)
            output_name = f"{safe_bank_name}_AEPS.docx"
            output_path = os.path.join(output_dir, output_name)
            
            success = self._fill_aeps_template(
                template_path, 
                bank_name, 
                transactions, 
                output_path,
                custom_subject,
                custom_message
            )
            
            if success:
                # Add DOCX
                generated_files.append({
                    'path': output_path,
                    'bank': bank_name,
                    'count': len(transactions),
                    'type': 'AEPS Withdrawal'
                })
                
                # Convert to PDF and add
                pdf_path = output_path.replace('.docx', '.pdf')
                if self._convert_to_pdf(output_path, output_dir) and os.path.exists(pdf_path):
                    generated_files.append({
                        'path': pdf_path,
                        'bank': f"{bank_name} (PDF)",
                        'count': len(transactions),
                        'type': 'AEPS Withdrawal'
                    })
        
        return generated_files
    
    def _fill_aeps_template(self, template_path, bank_name, transactions, output_path, custom_subject=None, custom_message=None):
        """Fill the AEPS template with transaction data"""
        try:
            doc = Document(template_path)
            
            # Replace placeholders in the document
            for paragraph in doc.paragraphs:
                # Replace bank name
                if 'Punjab National Bank' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Punjab National Bank', bank_name)
                
                # Replace custom subject if provided
                if custom_subject and '{{SUBJECT}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{{SUBJECT}}', custom_subject)
                
                # Replace custom message if provided
                if custom_message and '{{MESSAGE}}' in paragraph.text:
                    # Clear original paragraph text to remove placeholder without breaking XML
                    paragraph.text = paragraph.text.replace('{{MESSAGE}}', '')
                    
                    # Split message into lines and insert as new paragraphs with 0 spacing
                    lines = custom_message.split('\n')
                    for line in lines:
                        if not line.strip(): continue
                        new_p = paragraph.insert_paragraph_before(line)
                        new_p.paragraph_format.space_before = Pt(0)
                        new_p.paragraph_format.space_after = Pt(0)
                
                # Also check runs
                for run in paragraph.runs:
                    if 'Punjab National Bank' in run.text:
                        run.text = run.text.replace('Punjab National Bank', bank_name)
                    if custom_subject and '{{SUBJECT}}' in run.text:
                        run.text = run.text.replace('{{SUBJECT}}', custom_subject)

            
            # Find and populate the table
            if doc.tables:
                table = doc.tables[0]
                
                # Clear existing rows (keep header)
                for i in range(len(table.rows) - 1, 0, -1):
                    row = table.rows[i]
                    row._element.getparent().remove(row._element)
                
                # Add transaction rows
                for idx, txn in enumerate(transactions, 1):
                    row = table.add_row()
                    cells = row.cells
                    
                    # Template columns: Sl. No. | Account No. | Transaction ID | Withdrawal Date | Withdrawal Amount
                    if len(cells) >= 5:
                        cells[0].text = str(idx)
                        
                        # Column C: Account No./ (Wallet /PG/PA) Id
                        account_no = txn.get('Account No./ (Wallet /PG/PA) Id', '')
                        if pd.notna(account_no):
                            cells[1].text = str(int(account_no)) if isinstance(account_no, (int, float)) else str(account_no)
                        else:
                            cells[1].text = ''
                        
                        # Column D: Transaction Id / UTR Number
                        cells[2].text = str(txn.get('Transaction Id / UTR Number', '')) if pd.notna(txn.get('Transaction Id / UTR Number', '')) else ''
                        
                        # Column E: Withdrawal Date
                        cells[3].text = str(txn.get('Withdrawal Date', '')) if pd.notna(txn.get('Withdrawal Date', '')) else ''
                        
                        # Column F: Withdrawal Amount
                        amount = txn.get('Withdrawal Amount', '')
                        if pd.notna(amount):
                            cells[4].text = str(amount)
                        else:
                            cells[4].text = ''
            
            # Apply borders to the table
            if doc.tables:
                self.set_table_borders(doc.tables[0])
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error filling AEPS template: {e}")
            return False
